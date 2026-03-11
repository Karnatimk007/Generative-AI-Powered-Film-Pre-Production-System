import os
import io
import uuid
import base64
from datetime import datetime, timezone

from flask import Flask, render_template, request, jsonify, session, send_file, redirect
from flask_bcrypt import Bcrypt
from flask_login import (
    LoginManager, UserMixin,
    login_user, logout_user, login_required, current_user
)
from groq import Groq
from dotenv import load_dotenv

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors

# python-docx imports
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# MongoDB helpers
from db import create_user, find_user_by_email, get_user_by_id

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

load_dotenv(override=True)

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "scriptoria_secret_key_2024")

bcrypt      = Bcrypt(app)
login_mgr   = LoginManager(app)
login_mgr.login_view = "login_page"

# Server-side store — avoids Flask's 4 KB cookie limit for large AI outputs
_results_store: dict = {}

GROQ_MODEL  = "llama-3.3-70b-versatile"
# Fallback dummy key prevents server crash during platform build steps before env vars are wired
_groq_key   = os.getenv("GROQ_API_KEY", "dummy_key_for_builds")
groq_client = Groq(api_key=_groq_key)

# —— Startup: confirm HF key is loaded ——
_hf_key = os.getenv("HF_API_KEY", "").strip()
if _hf_key and _hf_key != "hf_your_token_here":
    print(f"[Scriptoria] ✅  HF_API_KEY loaded — nscale image gen ready: {_hf_key[:8]}...")
else:
    print("[Scriptoria] ⚠️  HF_API_KEY not set — shot image generation will fail.")

# ---------------------------------------------------------------------------
# Flask-Login user proxy (reads from MongoDB)
# ---------------------------------------------------------------------------

class MongoUser(UserMixin):
    """Thin wrapper so Flask-Login can work with MongoDB documents."""
    def __init__(self, doc: dict):
        self._id   = str(doc["_id"])
        self.name  = doc["name"]
        self.email = doc["email"]

    def get_id(self):
        return self._id


@login_mgr.user_loader
def load_user(user_id: str):
    doc = get_user_by_id(user_id)
    return MongoUser(doc) if doc else None


# ---------------------------------------------------------------------------
# Groq helper
# ---------------------------------------------------------------------------

def generate_with_groq(prompt):
    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=GROQ_MODEL,
            temperature=0.7,
            max_tokens=4096,
        )
        return chat_completion.choices[0].message.content.strip()
    except Exception as e:
        return f"[ERROR] Groq API error: {str(e)}"

# ---------------------------------------------------------------------------
# HuggingFace InferenceClient — nscale provider (SD XL Base)
# ---------------------------------------------------------------------------

_hf_client = None

def get_hf_client():
    """Return a cached InferenceClient using nscale as the provider."""
    global _hf_client
    if _hf_client is None:
        from huggingface_hub import InferenceClient
        api_key = os.getenv("HF_API_KEY", "").strip()
        if not api_key or api_key == "hf_your_token_here":
            raise RuntimeError("HF_API_KEY_NOT_SET")
        _hf_client = InferenceClient(
            provider="nscale",
            api_key=api_key,
        )
        print("[Scriptoria] ✅  nscale InferenceClient initialised.")
    return _hf_client

# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

def build_screenplay_prompt(story):
    return f"""You are a professional Hollywood screenwriter. Write a complete, industry-standard screenplay based on the following story idea.

STRICT FORMATTING RULES:
- Scene headings MUST be in ALL CAPS (e.g., INT. COFFEE SHOP - DAY)
- Character names before dialogue MUST be in ALL CAPS and centered
- Action lines in present tense, concise
- Include at least 3 scenes with dialogue
- Use proper screenplay structure: Setup, Confrontation, Resolution

STORY IDEA:
{story}

Write the full screenplay now:"""


def build_characters_prompt(story):
    return f"""You are a professional character designer and casting director. Based on the story below, create detailed character profiles for all major and supporting characters.

For EACH character provide:
- NAME (in ALL CAPS)
- AGE & PHYSICAL DESCRIPTION
- PERSONALITY TRAITS (3-5 bullet points)
- BACKSTORY (2-3 sentences)
- ROLE IN STORY
- SUGGESTED CASTING TYPE (e.g., "Mid-30s rugged male lead")
- CHARACTER ARC (how they change by the end)

STORY:
{story}

Generate all character profiles now:"""


def build_sound_design_prompt(story):
    return f"""You are an award-winning sound designer and composer. Create a comprehensive sound design plan for the following story.

Include:
1. OVERALL SONIC PALETTE (describe the emotional tone and audio world)
2. SCENE-BY-SCENE SOUND MAP:
   - Ambient sounds / room tone
   - Key sound effects (SFX)
   - Music cues (genre, tempo, instrumentation)
3. LEITMOTIFS: Recurring musical themes for main characters/concepts
4. SILENCE USAGE: Where silence is used for dramatic effect
5. TECHNICAL NOTES: Recommended recording techniques or sound libraries

STORY:
{story}

Create the full sound design plan now:"""


def build_script_breakdown_prompt(story):
    return f"""You are an experienced film production coordinator. Create a detailed script breakdown from the following story.

For EACH SCENE provide a breakdown table with:
- SCENE NUMBER & HEADING (INT/EXT, LOCATION, TIME)
- CAST REQUIREMENTS: Named characters appearing in scene
- PROPS: All physical objects needed
- LOCATIONS: Specific set or location requirements
- WARDROBE: Key costume notes
- SPECIAL REQUIREMENTS: Stunts, VFX, special equipment
- ESTIMATED SHOOT TIME: (in hours)

Also include at the end:
- TOTAL CAST LIST (all unique characters)
- MASTER PROPS LIST (all unique props)
- UNIQUE LOCATIONS LIST

STORY:
{story}

Generate the complete script breakdown now:"""


def build_shot_list_prompt(story):
    return f"""You are a cinematographer and director of photography (DP). Create a professional shot list for the following story.

For EACH SHOT provide:
- SHOT NUMBER
- SCENE REFERENCE
- SHOT TYPE: (e.g., ECU - Extreme Close Up, CU, MS, WS, EWS)
- CAMERA ANGLE: (e.g., Low Angle, High Angle, Bird's Eye, Dutch Tilt, Eye Level)
- CAMERA MOVEMENT: (e.g., Static, Pan, Tilt, Dolly, Handheld, Steadicam, Crane)
- LENS SUGGESTION: (e.g., 24mm wide, 35mm standard, 50mm normal, 85mm portrait, 135mm telephoto)
- LIGHTING MOOD: (e.g., Noir - high contrast, High-key - bright/flat, Golden Hour, Practical only, Motivated)
- DESCRIPTION: What the shot shows and its emotional purpose
- ESTIMATED DURATION: (in seconds)

STORY:
{story}

Generate the complete professional shot list now:"""


def build_image_prompt_from_shot(shot_description):
    return f"""You are a professional Hollywood cinematographer and AI prompt engineer.

Your task: Convert the following shot description into a highly detailed, cinematic AI image generation prompt.

Include ALL of these elements:
- Shot type (close-up, wide shot, aerial, extreme close-up, over-the-shoulder, etc.)
- Subject details (age, appearance, clothing, expression, body language)
- Action happening in the frame
- Environment/location with rich atmospheric detail
- Time of day (golden hour, blue hour, midnight, noon, etc.)
- Lighting style (Rembrandt, softbox, motivated practical, neon, candlelight, etc.)
- Camera angle (low angle, bird's eye, dutch tilt, eye level, etc.)
- Lens type and effect (24mm wide-angle, 35mm, 50mm standard, 85mm portrait bokeh, 135mm telephoto compression)
- Mood and emotion (tense, melancholic, euphoric, foreboding, romantic, etc.)
- Cinematic style reference (e.g., Blade Runner 2049, No Country for Old Men, La La Land, etc.)
- Color grading style (desaturated teal-orange, warm analog, cool monochrome, vivid Kodak)

End the prompt with: ultra realistic, 4K, high detail, film still, cinematic photography, anamorphic lens flare

SHOT DESCRIPTION:
{shot_description}

Output ONLY the final image prompt. No explanation, no preamble, no labels. Just the prompt itself."""

# ---------------------------------------------------------------------------
# Auth Routes
# ---------------------------------------------------------------------------

@app.route('/auth/register', methods=['POST'])
def auth_register():
    data     = request.get_json(force=True, silent=True) or {}
    name     = data.get('name', '').strip()
    email    = data.get('email', '').strip().lower()
    password = data.get('password', '')
    confirm  = data.get('confirm_password', '')

    if not name or not email or not password:
        return jsonify({'error': 'Name, email, and password are required.'}), 400
    if password != confirm:
        return jsonify({'error': 'Passwords do not match.'}), 400
    if len(password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters.'}), 400
    if find_user_by_email(email):
        return jsonify({'error': 'This email is already registered. Please log in.'}), 409

    pw_hash = bcrypt.generate_password_hash(password).decode('utf-8')
    user_id = create_user(name, email, pw_hash)

    from db import get_user_by_id
    doc  = get_user_by_id(user_id)
    user = MongoUser(doc)
    login_user(user, remember=True)
    session['sid'] = str(uuid.uuid4())

    return jsonify({'success': True, 'name': user.name, 'email': user.email}), 201


@app.route('/auth/login', methods=['POST'])
def auth_login():
    data     = request.get_json(force=True, silent=True) or {}
    email    = data.get('email', '').strip().lower()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({'error': 'Email and password are required.'}), 400

    doc = find_user_by_email(email)
    if not doc:
        print(f"[Auth] Login FAILED — no user found for: {email}")
        return jsonify({'error': 'No account found with that email. Please register first.'}), 401

    if not bcrypt.check_password_hash(doc['password_hash'], password):
        print(f"[Auth] Login FAILED — wrong password for: {email}")
        return jsonify({'error': 'Incorrect password. Please try again.'}), 401

    print(f"[Auth] Login OK — {email}")

    user = MongoUser(doc)
    login_user(user, remember=True)
    if 'sid' not in session:
        session['sid'] = str(uuid.uuid4())

    return jsonify({'success': True, 'name': user.name, 'email': user.email})


@app.route('/auth/logout', methods=['POST'])
@login_required
def auth_logout():
    logout_user()
    session.clear()
    return jsonify({'success': True})


@app.route('/auth/me', methods=['GET'])
def auth_me():
    if current_user.is_authenticated:
        return jsonify({'authenticated': True, 'name': current_user.name, 'email': current_user.email})
    return jsonify({'authenticated': False})

# ---------------------------------------------------------------------------
# Main routes
# ---------------------------------------------------------------------------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/login')
def login_page():
    """Dedicated login page."""
    if current_user.is_authenticated:
        return redirect('/')
    return render_template('login.html')


@app.route('/register')
def register_page():
    """Dedicated register page."""
    if current_user.is_authenticated:
        return redirect('/')
    return render_template('register.html')


@app.route('/generate', methods=['POST'])
@login_required
def generate():
    data  = request.get_json()
    story = data.get('story', '').strip()

    if not story:
        return jsonify({'error': 'Story input is required'}), 400

    if 'sid' not in session:
        session['sid'] = str(uuid.uuid4())
    sid = session['sid']

    session['story'] = story[:200]  # keep cookie small

    prompts = {
        'screenplay':       build_screenplay_prompt(story),
        'characters':       build_characters_prompt(story),
        'sound_design':     build_sound_design_prompt(story),
        'script_breakdown': build_script_breakdown_prompt(story),
        'shot_list':        build_shot_list_prompt(story),
    }

    results = {key: generate_with_groq(prompt) for key, prompt in prompts.items()}

    _results_store[sid] = results
    return jsonify({'success': True, 'results': results})


@app.route('/get_results', methods=['GET'])
def get_results():
    sid = session.get('sid')
    return jsonify(_results_store.get(sid, {}) if sid else {})


@app.route('/generate_shot_image', methods=['POST'])
@login_required
def generate_shot_image():
    """
    Shot description → Groq cinematic prompt → nscale SD XL image → base64 data-URI.
    """
    data      = request.get_json(force=True, silent=True) or {}
    shot_text = data.get('shot_description', '').strip()

    if not shot_text:
        return jsonify({'error': 'shot_description is required'}), 400

    # Step 1 — Groq crafts a rich Hollywood-DP image prompt
    image_prompt = generate_with_groq(build_image_prompt_from_shot(shot_text))
    if image_prompt.startswith('[ERROR]'):
        return jsonify({'error': image_prompt}), 500

    # Step 2 — Generate via nscale → SD XL Base (returns PIL.Image)
    try:
        client    = get_hf_client()
        pil_image = client.text_to_image(
            image_prompt[:900],   # keep within model token budget
            model="stabilityai/stable-diffusion-xl-base-1.0",
        )
    except RuntimeError as err:
        if "HF_API_KEY_NOT_SET" in str(err):
            return jsonify({
                'setup_required': True,
                'message': 'HF_API_KEY is missing. Add your Hugging Face token to .env as HF_API_KEY.',
            }), 200
        return jsonify({'error': str(err)}), 500
    except Exception as gen_err:
        return jsonify({'error': f'nscale image generation failed: {str(gen_err)}'}), 500

    # Step 3 — Convert PIL Image → base64 PNG data-URI
    buf = io.BytesIO()
    pil_image.save(buf, format='PNG')
    buf.seek(0)
    b64      = base64.b64encode(buf.read()).decode('utf-8')
    data_uri = f"data:image/png;base64,{b64}"

    return jsonify({
        'success':      True,
        'image_prompt': image_prompt,
        'image_url':    data_uri,
    })

# ---------------------------------------------------------------------------
# Export routes
# ---------------------------------------------------------------------------

@app.route('/export/<content_type>/<fmt>', methods=['POST'])
@login_required
def export(content_type, fmt):
    data    = request.get_json(force=True, silent=True) or {}
    content = data.get('content', '').strip()
    title   = content_type.replace('_', ' ').title()

    # Fallback: server-side store
    if not content:
        sid     = session.get('sid')
        results = _results_store.get(sid, {}) if sid else {}
        content = results.get(content_type, '')

    if not content:
        return jsonify({'error': 'No content found. Please generate first.'}), 404

    filename = f"scriptoria_{content_type}"

    # ---- TXT ----------------------------------------------------------------
    if fmt == 'txt':
        buf = io.BytesIO()
        buf.write(f"SCRIPTORIA — {title.upper()}\n".encode('utf-8'))
        buf.write(("=" * 60 + "\n\n").encode('utf-8'))
        buf.write(content.encode('utf-8'))
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                         as_attachment=True, download_name=f"{filename}.txt")

    # ---- PDF ----------------------------------------------------------------
    elif fmt == 'pdf':
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'ScriptoriaTitle', parent=styles['Title'],
            fontName='Times-Bold', fontSize=18,
            textColor=colors.HexColor('#1a1a2e'),
            spaceAfter=6, alignment=TA_CENTER
        )
        subtitle_style = ParagraphStyle(
            'ScriptoriaSubtitle', parent=styles['Normal'],
            fontName='Times-Italic', fontSize=11,
            textColor=colors.HexColor('#555555'),
            spaceAfter=20, alignment=TA_CENTER
        )
        body_style = ParagraphStyle(
            'ScriptoriaBody', parent=styles['Normal'],
            fontName='Times-Roman', fontSize=11,
            leading=16.5, spaceAfter=8, alignment=TA_LEFT
        )

        elements = [
            Paragraph("SCRIPTORIA", title_style),
            Paragraph(title, subtitle_style),
            Spacer(1, 0.2 * inch),
        ]
        for para in content.split('\n'):
            para = para.strip()
            if para:
                para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(para, body_style))
            else:
                elements.append(Spacer(1, 0.1 * inch))

        doc.build(elements)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name=f"{filename}.pdf")

    # ---- DOCX ---------------------------------------------------------------
    elif fmt == 'docx':
        buf = io.BytesIO()
        doc = Document()

        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        heading = doc.add_heading('SCRIPTORIA', level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Cambria'
            run.font.size = Pt(20)

        sub = doc.add_heading(title, level=1)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.name = 'Cambria'
            run.font.size = Pt(14)

        doc.add_paragraph()

        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Cambria'
                run.font.size = Pt(11)
            pPr     = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)

        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{filename}.docx"
        )

    return jsonify({'error': 'Invalid format. Use txt, pdf, or docx'}), 400

# ---------------------------------------------------------------------------

if __name__ == '__main__':
    app.run(debug=True, port=5000)
